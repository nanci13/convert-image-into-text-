import cv2
import pytesseract
from PIL import Image
from docx import Document

# If using Windows, set the Tesseract path:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def image_to_text(img_path):
    # Load image
    img = cv2.imread(img_path)

    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Apply thresholding for clarity
    gray = cv2.threshold(gray, 0, 255,
                         cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

    # Convert to PIL image
    pil_img = Image.fromarray(gray)

    # OCR extraction
    text = pytesseract.image_to_string(pil_img, lang="eng")
    return text

def save_to_doc(text, filename="output.docx"):
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"✅ Text saved to {filename}")

if __name__ == "__main__":
    img_file = "handwritten_sample.jpg"   # Replace with your image path
    extracted_text = image_to_text(img_file)
    print("Extracted Text:\n", extracted_text)

    # Save extracted text into Word document
    save_to_doc(extracted_text, "Extracted_Text.docx")
